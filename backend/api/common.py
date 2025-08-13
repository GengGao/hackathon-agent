from pathlib import Path
from typing import List

from fastapi import UploadFile
import io
import docx  # type: ignore
import pdfminer.high_level  # type: ignore
import pytesseract  # type: ignore
from PIL import Image, ImageOps  # type: ignore
import re
from html import unescape
from html.parser import HTMLParser

from rag import RuleRAG
import shutil
import os
from typing import Optional
import requests
from requests.exceptions import TooManyRedirects, RequestException


# Shared RAG instance (scoped per-session by callers)
rag = RuleRAG(Path(__file__).resolve().parents[1] / "docs" / "rules.txt", lazy=False)


MAX_FILE_BYTES = 10 * 1024 * 1024  # 10MB limit per file
ALLOWED_FILE_EXT = {".txt", ".md", ".pdf", ".docx", ".png", ".jpg", ".jpeg"}


def _configure_tesseract_binary() -> Optional[str]:
    """Best-effort configuration for the Tesseract binary on macOS/Homebrew.

    Returns the resolved command path if available, else None.
    """
    existing = shutil.which("tesseract")
    if existing:
        return existing
    for candidate in ("/opt/homebrew/bin/tesseract", "/usr/local/bin/tesseract"):
        if os.path.exists(candidate):
            pytesseract.pytesseract.tesseract_cmd = candidate  # type: ignore[attr-defined]
            return candidate
    return None

def extract_text_from_file(file: UploadFile) -> str:
    filename = file.filename or "uploaded_file"
    lower = filename.lower()
    ext = ("." + filename.split(".")[-1].lower()) if "." in filename else ""
    raw = file.file.read()
    if len(raw) > MAX_FILE_BYTES:
        return f"[File '{filename}' skipped: exceeds size limit]"
    # Explicitly communicate unsupported legacy .doc files
    if ext == ".doc":
        return (
            f"[File '{filename}' unsupported: .doc (Word 97-2003) is not supported. "
            "Please convert to .docx, .pdf, or upload an image for OCR.]"
        )
    if ext and ext not in ALLOWED_FILE_EXT:
        return f"[File '{filename}' skipped: extension not allowed]"
    try:
        if lower.endswith(".pdf"):
            return pdfminer.high_level.extract_text(io.BytesIO(raw))
        elif lower.endswith(".docx"):
            d = docx.Document(io.BytesIO(raw))
            parts: List[str] = []
            parts.extend(p.text for p in d.paragraphs if p.text)
            # Include table cell text which python-docx does not expose via paragraphs
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            return "\n".join(s.strip() for s in parts if s and s.strip())
        elif lower.endswith((".png", ".jpg", ".jpeg")):
            try:
                _configure_tesseract_binary()
                img = Image.open(io.BytesIO(raw))
                # Correct orientation based on EXIF and use a consistent mode for OCR
                img = ImageOps.exif_transpose(img)
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")
                text = pytesseract.image_to_string(img).strip()
                return text or f"[No text detected in image {filename}]"
            except Exception as e:  # pragma: no cover - best-effort OCR
                return f"[Image OCR failed for {filename}: {e}]"
        else:
            return raw.decode("utf-8", errors="ignore")
    except Exception as e:  # pragma: no cover - best-effort extraction
        return f"[Failed to process {filename}: {e}]"


class _BodyTextHTMLParser(HTMLParser):
    """Extract visible text from the HTML body, ignoring scripts/styles.

    Keeps simple structure hints by injecting newlines around common block-level tags.
    """

    def __init__(self) -> None:
        super().__init__()
        self.in_body: bool = False
        self.suppressed_depth: int = 0  # inside <script>/<style>/<noscript>
        self.chunks: List[str] = []
        self.seen_body: bool = False

    def handle_starttag(self, tag: str, attrs):  # type: ignore[override]
        tag_lower = tag.lower()
        if tag_lower == "body":
            self.in_body = True
            self.seen_body = True
            return
        if not self.in_body:
            return
        if tag_lower in ("script", "style", "noscript"):
            self.suppressed_depth += 1
            return
        if tag_lower in ("br",):
            self.chunks.append("\n")
            return
        if tag_lower in (
            "p",
            "div",
            "section",
            "header",
            "footer",
            "article",
            "li",
            "ul",
            "ol",
            "table",
            "thead",
            "tbody",
            "tr",
            "td",
            "th",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "pre",
        ):
            self.chunks.append("\n")

    def handle_endtag(self, tag: str):  # type: ignore[override]
        tag_lower = tag.lower()
        if tag_lower == "body":
            self.in_body = False
            return
        if not self.in_body:
            return
        if tag_lower in ("script", "style", "noscript"):
            if self.suppressed_depth > 0:
                self.suppressed_depth -= 1
            return
        if tag_lower in (
            "p",
            "div",
            "section",
            "header",
            "footer",
            "article",
            "li",
            "ul",
            "ol",
            "table",
            "thead",
            "tbody",
            "tr",
            "td",
            "th",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "pre",
        ):
            self.chunks.append("\n")

    def handle_data(self, data: str):  # type: ignore[override]
        if not self.in_body or self.suppressed_depth > 0:
            return
        if not data or not data.strip():
            return
        self.chunks.append(data)


def extract_visible_text_from_html(html_text: str) -> str:
    """Return human-visible text from the <body> of an HTML document.

    Falls back to stripping head/script/style tags if no explicit <body> exists.
    """
    if not html_text:
        return ""

    parser = _BodyTextHTMLParser()
    parser.feed(html_text)
    body_text = "".join(parser.chunks).strip()

    if not parser.seen_body:
        # Fallback: strip obvious non-visible sections and tags
        cleaned = re.sub(r"(?is)<(script|style|noscript)[^>]*>.*?</\\1>", " ", html_text)
        cleaned = re.sub(r"(?is)<head[^>]*>.*?</head>", " ", cleaned)
        cleaned = re.sub(r"(?is)<[^>]+>", " ", cleaned)
        body_text = cleaned

    text = unescape(body_text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n[ \t]*\n+", "\n\n", text)
    return text.strip()


def build_url_block(url: str, *, timeout: int = 5, max_bytes: int = 100_000, max_redirects: int = 3) -> str:
    # Allow only human-readable text types
    def _is_allowed_mime(ctype_raw: str) -> bool:
        if not ctype_raw:
            return False
        ctype = ctype_raw.split(";", 1)[0].strip().lower()
        if ctype.startswith("text/"):
            return True
        # Allow XHTML explicitly (served as application/xhtml+xml)
        if ctype in {"application/xhtml+xml"}:
            return True
        return False

    session = requests.Session()
    session.max_redirects = max_redirects

    # HEAD size/mime guard
    try:
        head_resp = session.head(url, timeout=timeout, allow_redirects=True)
        head_ctype = head_resp.headers.get("Content-Type", "")
        if not _is_allowed_mime(head_ctype):
            return f"[URL:{url}]\n[Blocked non-text content-type {head_ctype}]\n[/URL]"
        clen = head_resp.headers.get("Content-Length")
        if clen is not None:
            try:
                size_int = int(clen)
                if size_int > max_bytes:
                    return f"[URL:{url}]\n[Blocked: content-length {size_int} exceeds limit {max_bytes}]\n[/URL]"
            except ValueError:
                # Ignore invalid content-length and proceed to GET with streaming limits
                pass
    except TooManyRedirects:
        return f"[URL_FETCH_FAILED:{url}]\nError: too many redirects (> {max_redirects})"
    except Exception:
        # If HEAD fails (405 or network), proceed to GET with streaming safeguards
        pass

    # GET with streaming and hard byte cap; avoid buffering full response
    resp = None
    try:
        resp = session.get(url, timeout=timeout, stream=True, allow_redirects=True)
        ctype = resp.headers.get("Content-Type", "")
        if not _is_allowed_mime(ctype):
            return f"[URL:{url}]\n[Blocked non-text content-type {ctype}]\n[/URL]"

        total = 0
        chunks = []
        for chunk in resp.iter_content(chunk_size=8192):
            if not chunk:
                continue
            chunks.append(chunk)
            total += len(chunk)
            if total >= max_bytes:
                break

        content_bytes = b"".join(chunks)
        is_truncated = total >= max_bytes or False
        if len(content_bytes) > max_bytes:
            content_bytes = content_bytes[:max_bytes]
            is_truncated = True

        lower_ctype = ctype.split(";", 1)[0].strip().lower()
        if "html" in lower_ctype:
            try:
                html_text = content_bytes.decode("utf-8", errors="ignore")
                visible = extract_visible_text_from_html(html_text)
                snippet = visible + ("\n[Truncated]" if is_truncated else "")
            except Exception as e:  # pragma: no cover - best-effort HTML parsing
                snippet = f"[Failed HTML parse: {e}]"
        else:
            snippet = content_bytes.decode("utf-8", errors="ignore")
            if is_truncated:
                snippet += "\n[Truncated]"

        return f"[URL:{url}]\n{snippet}\n[/URL]"
    except TooManyRedirects:
        return f"[URL_FETCH_FAILED:{url}]\nError: too many redirects (> {max_redirects})"
    except RequestException as e:
        return f"[URL_FETCH_FAILED:{url}]\nError: {e}"
    except Exception as e:
        return f"[URL_FETCH_FAILED:{url}]\nError: {e}"
    finally:
        try:
            if resp is not None:
                resp.close()
        finally:
            session.close()


def get_generate_stream():
    """Return the streaming function used by chat routes.

    Tests monkeypatch `router.generate_stream`; to preserve backward compatibility
    after refactor, fetch from `router` module when available, otherwise fall back
    to the implementation in `llm`.
    """
    try:
        import router as router_module  # type: ignore

        gs = getattr(router_module, "generate_stream", None)
        if callable(gs):
            return gs
    except Exception:
        pass
    from llm import generate_stream as real_generate_stream  # lazy import to avoid cycles

    return real_generate_stream


